# -*- coding: utf-8 -*-
"""
Created on Sat Feb  8 14:09:35 2020

@author: Okamoto-DRIHA001
"""

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SimpleDocxService:
    
    def __init__(self):
        # コンストラクタ
        self.document = Document()
        self.latest_run = None

    def set_normal_font(self, name, size):
        # フォントの設定
        font = self.document.styles['Normal'].font
        font.name = name
        font.size = Pt(size)

    def add_head(self, text, lv):
        # 見出しの設定
        self.document.add_heading(text, level=lv)

    def open_text(self):
        # テキスト追加開始
        self.paragraph = self.document.add_paragraph()

    def close_text(self):
        # テキスト追加終了
        return # 現状では特に処理はなし

    def get_unicode_text(self, text, src_code):
        # python-docxで扱えるようにunicodeに変換
        return unicode(text, src_code)

    def adjust_return_code(self, text):
        # テキストファイルのデータをそのままaddすると改行が面倒なことになるので、それを削除
        text = text.replace("\n", "")
        text = text.replace("\r", "")
        return text

    def add_text(self, text):
        # テキスト追加
        self.latest_run = self.paragraph.add_run(text)

    def add_text_italic(self, text):
        # テキスト追加（イタリックに）
        self.paragraph.add_run(text).italic = True

    def add_text_bold(self, text):
        # テキスト追加（強調）
        self.paragraph.add_run(text).bold = True

    def add_text_color(self, text, r, g, b):
        # 文字に色をつける
        self.paragraph.add_run(text).font.color.rgb = RGBColor(r, g, b)

    def paragraph_alignment_center(self):
        # テキストを中央に配置する
        paragraph_format = self.paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def paragraph_alignment_left(self):
        # テキストを左に配置する
        paragraph_format = self.paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    def paragraph_alignment_right(self):
        # テキストを右に配置する
        paragraph_format = self.paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_picture(self, filename, inch):
        # 図を挿入する
        self.document.add_picture(filename, width=Inches(inch))

    def add_table(self,rows,cols):
        # 表を挿入する
        self.table = self.document.add_table(rows, cols)
        
    def add_value_to_table(self,rows,cols,value):
        # 表に値を挿入する
        row_cells = self.table.rows[rows].cells
        row_cells[cols].text = value
        
    def columns_width(self, cols, width):
        # 表の列幅を変更する
        for cell in self.table.columns[cols].cells:
            cell.width = Cm(width)
    
    def add_page_break(self):
        # 改ページを行う
        self.document.add_page_break()

    def save(self, name):
        # docxファイルとして出力。
        self.document.save(name)